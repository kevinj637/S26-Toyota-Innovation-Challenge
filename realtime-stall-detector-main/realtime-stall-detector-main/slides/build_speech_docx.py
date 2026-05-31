"""Build the speaker script as a Word doc (slides/SPEECH.docx) with python-docx.

Same words as SPEECH.md / the deck's notes, formatted for reading aloud: per-slide
spoken lines with **bold** emphasis and *italic* stage directions, a timing table,
rubric coverage, and a Q&A crib sheet.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
HERE = os.path.dirname(os.path.abspath(__file__))
GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def rich(p, text, size=13):
    """Render inline **bold** and *italic* markers into runs."""
    pat = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*')
    pos = 0
    for m in pat.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()]); r.font.size = Pt(size)
        if m.group(1) is not None:
            r = p.add_run(m.group(1)); r.bold = True; r.font.size = Pt(size)
        else:
            r = p.add_run(m.group(2)); r.italic = True; r.font.size = Pt(size); r.font.color.rgb = GRAY
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); r.font.size = Pt(size)


doc.add_heading("Pitch Speech — Real-Time Motor Stall Detection", level=0)
sub = doc.add_paragraph(); r = sub.add_run("≈ 2:45  ·  Toyota Innovation Challenge — Fault Prediction")
r.italic = True; r.font.color.rgb = GRAY
rich(doc.add_paragraph(),
     "Read top to bottom. **[SLIDE n]** = advance the slide.  **Bold** = emphasis.  *[brackets]* = stage action.", 11)

slides = [
    ("SLIDE 1 · Title",
     "Unexpected machine downtime costs manufacturers **millions a year**. The earliest warning a motor gives "
     "before it fails is its own **current** — so we built a system that reads that current in real time and "
     "catches a stall the instant it happens."),
    ("SLIDE 2 · The trap",
     "Here's what makes this hard. When a motor stalls, current spikes — but current **also** spikes every single "
     "time the motor starts: the inrush. A naive 'current-too-high' alarm would cry wolf on every power-on. The "
     "real signal isn't *high* current — it's high current that **stays** high. A stall is sustained; a start-up "
     "just decays. Telling those two apart is the heart of our project."),
    ("SLIDE 3 · How it works",
     "Under the hood: we slice the current into short windows, pull out **six features** — level, slope, and "
     "whether it's still rising or holding — and feed a small **decision tree**. We chose a tree on purpose: it's "
     "**interpretable**, and tiny enough to run on a chip. And the algorithm itself discovered the key rule — "
     "that **steadiness** is what separates a real stall from a start-up."),
    ("SLIDE 4 · Live demo",
     "*[switch to the live dashboard / the bench]*  Let me show you. Running — **green**. Watch when I power-cycle "
     "it — *[power-cycle the motor]* — the current spikes, but it stays **green**. No false alarm. Now I actually "
     "stall it — *[hold the shaft]* — **RED**, in under a tenth of a second. Release — back to green. Start-up "
     "ignored, real stall caught instantly."),
    ("SLIDE 5 · Results",
     "On our collected data it cleanly separates the three states. Stall is **perfectly classified** — never "
     "missed — and a start-up is **never** predicted as a stall. That's the property that matters. The whole "
     "pipeline has **eighteen automated tests** and a physics simulator, so we validated it before we ever "
     "touched hardware."),
    ("SLIDE 6 · Edge ML",
     "And it doesn't need a laptop. We compiled that **same** decision tree into C and flashed it onto the Arduino "
     "itself. The board runs the inference **on-chip** and lights this LED on a stall. *[point to the board's LED]* "
     "Pull the laptop, power it from a charger — it still works. Detection at the edge, on a microcontroller — "
     "exactly where a factory sensor lives."),
    ("SLIDE 7 · Scale",
     "And the method isn't locked to our bench: the **exact same feature code** runs unchanged on Toyota's real "
     "eight-DOF arm telemetry, across the whole fleet — and it surfaces the most mechanically-stressed joints, the "
     "shoulder and elbow, exactly where the load is. With labeled fault data, the same approach scales to "
     "**production predictive maintenance**."),
    ("SLIDE 8 · Safety, Human, Close",
     "It's human-centered: it catches failures before they cascade — protecting equipment, preventing unplanned "
     "line stops, and keeping workers out of harm's way. **From a three-dollar motor to a robot fleet — same "
     "physics, same code.** We turn a motor's current into an early-warning system. **Thank you.**"),
]
for title, line in slides:
    hp = doc.add_paragraph(); hr = hp.add_run(title); hr.bold = True; hr.font.size = Pt(13); hr.font.color.rgb = BLUE
    hp.paragraph_format.space_before = Pt(10); hp.paragraph_format.space_after = Pt(2)
    bp = doc.add_paragraph(); bp.paragraph_format.space_after = Pt(6); rich(bp, line, 13)

doc.add_paragraph()
tr = doc.add_paragraph().add_run("Timing (~2:45)"); tr.bold = True
tbl = doc.add_table(rows=2, cols=8); tbl.style = "Light Grid Accent 1"
for i, c in enumerate(["S1", "S2", "S3", "S4 demo", "S5", "S6", "S7", "S8"]):
    tbl.rows[0].cells[i].text = c
for i, c in enumerate(["0:15", "0:30", "0:30", "0:45", "0:20", "0:20", "0:20", "0:15"]):
    tbl.rows[1].cells[i].text = c

rp = doc.add_paragraph(); rp.add_run("Rubric coverage: ").bold = True
rp.add_run("Ideation → inrush-vs-stall insight (S2) · Execution → live demo + tests (S4/S5) · "
           "Safety & Human → S8 · Presentation → interpretable tree + clean visuals.")

qh = doc.add_paragraph(); qhr = qh.add_run("If asked (Q&A)"); qhr.bold = True; qhr.font.size = Pt(13)
qh.paragraph_format.space_before = Pt(12)
qa = [
    ("“Did you reuse the small-motor model on the arm?”",
     "No. Only the feature-extraction code transfers; each deployment trains its own model. We do not claim arm "
     "stall detection — that data is unlabeled."),
    ("“Why a decision tree, not a neural net?”",
     "Interpretable, fits on a chip, and it's already perfect on the critical class. We can show the exact "
     "learned rule."),
    ("“Different motor / generalization?”",
     "Recollect a few minutes, retrain in seconds; the features are threshold-free, and we train and infer on the "
     "same scale."),
    ("“Are the Toyota red cells faults?”",
     "No — highest mechanical-stress windows, not confirmed faults. A real jam would look the same and be caught "
     "the same way."),
    ("“How is this different from a current relay?”",
     "A relay only sees a threshold — it false-alarms on every start-up. We distinguish inrush from stall."),
]
for q, a in qa:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.add_run(q + "  ").bold = True; p.add_run("→ " + a)

out = os.path.join(HERE, "SPEECH.docx")
doc.save(out)
print("saved", out)
